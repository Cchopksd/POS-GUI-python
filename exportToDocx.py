from unittest import result
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from tkinter import messagebox
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
import calendar
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import *

def createDocxFile(result_list,report_type,report_date,total_price,total_amount) :

    messagebox.showinfo("ระบบ:","พิมพ์รายงานสำเร็จ")

    date_txt = report_date.replace('/', '-')

    document = Document()
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    if report_type == "รายงานสรุปยอดขายสินค้า" :
        sell_result = document.add_paragraph('รายงานสรุปยอดขายสินค้า\nณ วันที่ : '+str(report_date))
        sell_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'เลขที่ออเดอร์'
        hdr_cells[1].text = 'ผู้ออกออเดอร์'
        hdr_cells[2].text = 'รายละเอียดออเดอร์'
        hdr_cells[3].text = 'ราคารวมออเดอร์'
        hdr_cells[4].text = 'เงินที่รับมา'
        hdr_cells[5].text = 'เงินทอน'
        for orderId, orderSeller, orderDesc, orderPrice, orderGetMoney, orderChange in result_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(orderId)
            row_cells[1].text = str(orderSeller)
            row_cells[2].text = str(orderDesc)
            row_cells[3].text = str(orderPrice) + " บาท"
            row_cells[4].text = str(orderGetMoney)+ " บาท"
            row_cells[5].text = str(orderChange)+ " บาท"

        total_result = document.add_paragraph('\n\nราคารวม                           '+str(total_price)+'                           บาท')
        total_result.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph('\nรายงานสรุปยอดขายสินค้า\nวันที่ : '+str(report_date))

    elif report_type == "รายงานการสั่งซื้อวัตถุดิบ" :
        sell_result = document.add_paragraph('รายงานการสั่งซื้อวัตถุดิบ\nณ วันที่ : '+str(report_date))
        sell_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'เลขที่สั่งซื้อวัตถุดิบ'
        hdr_cells[1].text = 'ผู้สั่งซื้อวัตถุดิบ'
        hdr_cells[2].text = 'รายละเอียดการสั่งซื้อ'
        hdr_cells[3].text = 'จำนวนการสั่งซื้อ'
        hdr_cells[4].text = 'ราคาการสั่งซื้อ'
        hdr_cells[5].text = 'บริษัทที่ทำการสั่งซื้อ'
        for stapleId, stapleBuyer, stapleDesc, stapleAmount, stapleUnit, staplePrice, stapleSource in result_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(stapleId)
            row_cells[1].text = str(stapleBuyer)
            row_cells[2].text = str(stapleDesc)
            row_cells[3].text = str(stapleAmount) + " "+ str(stapleUnit)
            row_cells[4].text = str(staplePrice) + " บาท"
            row_cells[5].text = str(stapleSource)

        document.add_paragraph('\n\n                                                                                                                                                                 จำนวนครั้งที่สั่งซื้อ                             '+str(total_amount)+'                                ครั้ง')
        total_result = document.add_paragraph('ราคาที่สั่งซื้อทั้งหมด                           '+str(total_price)+'                           บาท')
        total_result.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph('\nรายงานการสั่งซื้อวัตถุดิบ\nวันที่ : '+str(report_date))
    
    elif report_type == "รายงานการเบิกวัตถุดิบ" :
        sell_result = document.add_paragraph('รายงานการเบิกวัตถุดิบ\nณ วันที่ : '+str(report_date))
        sell_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'เลขที่การเบิกวัตถุดิบ'
        hdr_cells[1].text = 'ผู้ทำการเบิกวัตถุดิบ'
        hdr_cells[2].text = 'รายละเอียดการเบิกวัตถุดิบ'
        for req_Id, req_person, req_list in result_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(req_Id)
            row_cells[1].text = str(req_person)
            row_cells[2].text = str(req_list)

        total_result = document.add_paragraph('\n\nจำนวนการเบิกทั้งหมด                           '+str(total_amount)+'                           ครั้ง')
        total_result.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph('\nรายงานการเบิกวัตถุดิบ\nวันที่ : '+str(report_date))

    # result_list.clear()
    # print(result_list)
    file_name = report_type + "(" + str(date_txt) + ")"
    document.save(file_name + '.docx')
